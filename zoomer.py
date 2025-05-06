from flask import Flask, render_template, request, jsonify
from openai import OpenAI
import datetime
from docx import Document
import os
from dotenv import load_dotenv
import numpy as np
import faiss

app = Flask(__name__)

load_dotenv()  # Load environment variables from .env file

# Initialize NetMind API
client = OpenAI(
    base_url="https://api.netmind.ai/inference-api/openai/v1",
    api_key=os.getenv("OPEN_API_KEY")
)

# ===== Load and Embed Zoomer Africa FAQ =====
doc = Document("zoomer_faqs.docx")
faq_chunks = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

def get_embedding(text):
    response = client.embeddings.create(
        model="BAAI/bge-m3",  # ✅ Use correct model from NetMind
        input=text
    )
    return response.data[0].embedding

# Build FAISS index
# sample_emb = get_embedding("sample")  # to determine dimension
# embedding_dim = len(sample_emb)
# index = faiss.IndexFlatL2(embedding_dim)

# faq_embeddings = []
# for chunk in faq_chunks:
#     emb = get_embedding(chunk)
#     faq_embeddings.append(emb)
#     index.add(np.array([emb]).astype("float32"))

# ========== AfroZoomer Assistant Class ==========
class AfroZoomerAssistant:
    def __init__(self):
        self.name = "AfroZoomer"
        self.conversation_history = []

        # Build FAISS index inside the assistant
        self.faq_chunks = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        sample_emb = get_embedding("sample")
        self.embedding_dim = len(sample_emb)
        self.index = faiss.IndexFlatL2(self.embedding_dim)

        for chunk in self.faq_chunks:
            emb = get_embedding(chunk)
            self.index.add(np.array([emb]).astype("float32"))

    def get_contextual_faq(self, user_prompt):
        emb = get_embedding(user_prompt)
        emb_np = np.array([emb]).astype("float32")
        _, I = self.index.search(emb_np, k=3)  # top 3 FAQ matches
        relevant_faqs = [faq_chunks[i] for i in I[0]]
        return "\n".join(relevant_faqs)

    def get_response(self, user_input):
        context = self.get_contextual_faq(user_input)
        messages = [
            {"role": "system", "content": "You are AfroZoomer, an assistant who answers questions about Zoomer Africa."},
            {"role": "user", "content": f"Context:\n{context}\n\nQuestion: {user_input}"}
        ]
        response = client.chat.completions.create(
            model="Qwen/Qwen3-8B",
            messages=messages,
            max_tokens=4096,
            temperature=0.6,
            top_p=0.9,
        )
        return response.choices[0].message.content.strip()

assistant = AfroZoomerAssistant()

# ========== Flask Routes ==========
@app.route('/')
def home():
    return render_template("zoomer.html")

@app.route('/ask', methods=['POST'])
def ask():
    user_input = request.form['user_input']
    response = assistant.get_response(user_input)
    return jsonify({'response': response})

if __name__ == "__main__":
    app.run(debug=True, port=8000)
